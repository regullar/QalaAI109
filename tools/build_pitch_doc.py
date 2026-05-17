from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = "deliverables/Qala_AI_3_minute_pitch_timeline.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_tag(paragraph, text, fill="EAF2FF"):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(18, 70, 135)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p


def set_table_borders(table, color="D7DEE8"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_timeline_table(doc):
    rows = [
        (
            "0:00-0:20",
            "Главная страница",
            "Проблема",
            "В Шымкенте уже есть I-Shymkent 109, куда жители обращаются по городским проблемам. По данным Kazinform от 23 апреля 2026 года, с начала 2026 года в контакт-центр поступило около 111 тысяч обращений, операторы принимают 1000-1500 звонков в день, а более 28 тысяч звонков касались отключений электроэнергии. Проблема не в том, что жители не пишут. Проблема в хаотичном потоке: оператору нужно быстро понять, что случилось, где, насколько срочно и кому передать.",
            "Жюри понимает: проблема реальная, локальная, измеримая.",
        ),
        (
            "0:20-0:45",
            "Hero: Qala AI + блок 'Главная идея'",
            "Решение",
            "Qala AI - это не просто сайт жалоб и не 'ChatGPT сортирует обращения'. Это AI-диспетчерский слой для городских обращений Шымкента. Он превращает сообщение жителя обычным языком в структурированную заявку: категория, район, приоритет, риск-факторы, ответственная служба и текст обращения. Главное отличие: система видит не только отдельные заявки, а повторяющиеся проблемные зоны.",
            "Сразу закрепить позиционирование: Smart City диспетчер, а не форма.",
        ),
        (
            "0:45-1:30",
            "/report",
            "AI-диспетчеризация",
            "Покажу на примере. Житель пишет: 'В мкр Нурсат возле школы вечером не горят фонари, дети идут домой в темноте'. Пользователь не обязан знать категорию или службу. Система сама анализирует текст, определяет проблему как уличное освещение, учитывает риск 'дети + школа + темное место', выставляет высокий приоритет и предлагает ответственную службу.",
            "Показать живой путь от сырого текста до готовой заявки.",
        ),
        (
            "1:30-1:55",
            "AI Preview на /report",
            "Результат анализа",
            "Для оператора это уже не сырой текст из WhatsApp или звонка. Это управляемая карточка: заголовок, краткое описание, район, категория, срочность, риск-факторы, служба и официальный текст обращения. Если внешний AI недоступен, есть fallback classifier, поэтому MVP не ломается во время демо.",
            "Подчеркнуть надежность и практичность.",
        ),
        (
            "1:55-2:20",
            "/map",
            "Карта и кластеры",
            "Теперь ключевая фича. Если 10 жителей разными словами пишут про одну проблему в Нурсате, для города это не 10 отдельных жалоб. Это одна проблемная зона. Qala AI объединяет похожие обращения по району, категории, координатам и смыслу, а затем показывает горячие зоны на карте Шымкента через 2GIS.",
            "Показать Smart City эффект: город видит зоны риска.",
        ),
        (
            "2:20-2:40",
            "Telegram-бот или слайд с командами",
            "Канал сбора",
            "Отдельно мы добавили Telegram-бота как канал входящих обращений. Его можно подключить к чату ЖК или двора: команда /qala_start сохраняет адрес или геолокацию, /qala_collect включает окно сбора сообщений, бот фильтрует короткий шум вроде 'ок' и '+1', группирует похожие сообщения и отправляет их в Qala AI как заявки из источника Telegram Demo.",
            "Доказать, что это не только веб-форма, а входящий поток из реального канала.",
        ),
        (
            "2:40-2:52",
            "/admin и /admin/analytics",
            "Панель оператора",
            "У оператора есть очередь заявок со статусами, приоритетами и похожими обращениями. У акимата есть аналитика: какие районы перегружены, какие категории растут, где повторяются проблемы.",
            "Доказать, что MVP полезен и жителю, и оператору, и акимату.",
        ),
        (
            "2:52-3:00",
            "Главная или аналитика",
            "Финал",
            "Мы не заменяем I-Shymkent 109. Мы показываем AI-слой, который снижает нагрузку на операторов, быстрее находит срочные случаи и показывает городу реальные проблемные зоны.",
            "Завершить практической ценностью и масштабированием.",
        ),
    ]

    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Cm(2.2), Cm(3.8), Cm(3.2), Cm(11.6), Cm(5.0)]
    headers = ["Время", "Что показать", "Задача", "Слова спикера", "Что должно остаться у жюри"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        set_cell_shading(cell, "1F4E79")
        set_cell_text(cell, header, bold=True, color=(255, 255, 255))

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = widths[i]
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_text(cells[i], value)
        set_cell_shading(cells[0], "EEF5FF")
        set_cell_shading(cells[2], "F7FAFC")

    set_table_borders(table)
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].paragraph_format.space_after = Pt(5)
    styles["Title"].font.name = "Arial"
    styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Title"].font.size = Pt(22)
    styles["Title"].font.bold = True
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Heading 1"].font.size = Pt(14)
    styles["Heading 1"].font.bold = True
    styles["Heading 1"].font.color.rgb = RGBColor(31, 78, 121)
    styles["Heading 2"].font.name = "Arial"
    styles["Heading 2"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    styles["Heading 2"].font.size = Pt(12)
    styles["Heading 2"].font.bold = True
    styles["Heading 2"].font.color.rgb = RGBColor(31, 78, 121)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Qala AI: 3-минутный спич для защиты")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("AI-диспетчер городских обращений для Шымкента").bold = True
    subtitle.paragraph_format.space_after = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Главная мысль: ").bold = True
    p.add_run("обычные системы видят 100 отдельных жалоб, Qala AI показывает реальные проблемные зоны города.")

    doc.add_heading("1. Позиционирование", level=1)
    add_bullet(doc, "Не говорить: 'мы сделали сайт жалоб'.")
    add_bullet(doc, "Не говорить: 'ChatGPT сортирует обращения'.")
    add_bullet(doc, "Говорить: 'Qala AI - AI-диспетчерский слой для I-Shymkent 109 и городских обращений Шымкента'.")
    add_bullet(doc, "Название на сайте оставляем Qala AI. В речи можно уточнять: 'по сути, это Shymkent 109 AI Hub'.")

    callout = doc.add_table(rows=1, cols=1)
    callout.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = callout.cell(0, 0)
    set_cell_shading(cell, "EAF2FF")
    set_cell_text(
        cell,
        "Одна фраза проекта: Qala AI превращает хаотичные обращения жителей в структурированные заявки, объединяет повторяющиеся жалобы в проблемные зоны и показывает их на карте Шымкента.",
        bold=True,
    )
    set_table_borders(callout, "B7C9E2")

    doc.add_heading("2. Таймлайн на 3 минуты", level=1)
    add_timeline_table(doc)

    doc.add_page_break()
    doc.add_heading("3. Демо-маршрут: что открыть заранее", level=1)
    for item in [
        "Вкладка 1: главная страница `/` - открыть на hero Qala AI.",
        "Вкладка 2: `/report` - заранее пройти Clerk sign-in, чтобы не показывать логин жюри.",
        "Вкладка 3: `/map` - заранее проверить, что есть seed-данные и видны кластеры.",
        "Вкладка 4: `/admin` - очередь заявок.",
        "Вкладка 5: `/admin/analytics` - аналитика по районам, категориям и кластерам.",
        "Опционально: открыть Telegram-чат с ботом или отдельный слайд со схемой команд `/qala_start`, `/qala_collect 30m`, `/qala_collect_stop`.",
    ]:
        add_number(doc, item)

    doc.add_heading("4. Текст для живого примера", level=1)
    example = doc.add_table(rows=2, cols=2)
    example.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(["Поле", "Что использовать"]):
        set_cell_shading(example.rows[0].cells[i], "1F4E79")
        set_cell_text(example.rows[0].cells[i], header, bold=True, color=(255, 255, 255))
    set_cell_text(example.rows[1].cells[0], "Текст обращения", bold=True)
    set_cell_text(
        example.rows[1].cells[1],
        "В мкр Нурсат возле школы вечером не горят фонари, дети идут домой в темноте",
    )
    set_table_borders(example)

    doc.add_heading("5. Самые сильные фразы", level=1)
    strong_phrases = [
        "Это не форма жалоб, а AI-диспетчерский слой для города.",
        "Житель пишет обычным языком, оператор получает структурированную заявку.",
        "Для города 10 похожих жалоб - это не 10 отдельных проблем, а одна проблемная зона.",
        "Qala AI помогает видеть не шум обращений, а карту городских рисков.",
        "Telegram-бот показывает, что система может собирать обращения не только из веб-формы, но и из живых чатов дворов и ЖК.",
        "Мы не заменяем I-Shymkent 109, мы показываем AI-слой, который может ускорить классификацию, приоритизацию и контроль.",
    ]
    for phrase in strong_phrases:
        add_bullet(doc, phrase)

    doc.add_heading("6. Если жюри задаст вопросы", level=1)
    qa = [
        (
            "Это официальный сервис 109?",
            "Нет. Это демо Smart City MVP. Мы осознанно показываем дисклеймер, что проект не является официальным сервисом акимата или I-Shymkent 109.",
        ),
        (
            "Что делает AI?",
            "AI классифицирует обращение, выделяет район, приоритет, риск-факторы, ответственную службу и формирует официальный текст обращения.",
        ),
        (
            "Что если AI недоступен?",
            "Есть fallback classifier. Базовая диспетчеризация продолжает работать, поэтому демо не зависит полностью от внешнего AI.",
        ),
        (
            "Почему это Smart City?",
            "Потому что система превращает обращения жителей в городские данные: карту проблем, повторяющиеся зоны, аналитику по районам и приоритетам.",
        ),
        (
            "Зачем Telegram-бот, если есть сайт?",
            "Сайт удобен для индивидуального обращения, а Telegram-бот нужен для реального дворового потока: жители уже пишут в чатах ЖК, а бот превращает этот шум в заявки.",
        ),
        (
            "Что именно умеет бот?",
            "Он подключает чат, сохраняет адрес или геолокацию, включает окно сбора сообщений, фильтрует шум, группирует похожие сообщения и отправляет заявки в Qala AI.",
        ),
        (
            "Как масштабировать?",
            "Подключить источники 109, WhatsApp, Telegram, мобильное приложение, хранить обращения в единой базе и строить аналитику по районам и службам. Telegram-бот уже показывает один такой канал.",
        ),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(["Вопрос", "Короткий ответ"]):
        set_cell_shading(table.rows[0].cells[i], "1F4E79")
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=(255, 255, 255))
    for q, a in qa:
        cells = table.add_row().cells
        set_cell_text(cells[0], q, bold=True)
        set_cell_text(cells[1], a)
    set_table_borders(table)

    doc.add_heading("7. Источник цифр", level=1)
    p = doc.add_paragraph()
    p.add_run("Kazinform, 23 апреля 2026: ").bold = True
    p.add_run(
        "с начала 2026 года в I-Shymkent 109 поступило около 111 тысяч обращений; операторы принимают 1000-1500 звонков в день; более 28 тысяч звонков касались отключений электроэнергии; около 6 тысяч обращений поступило по дворовому и уличному освещению."
    )
    p2 = doc.add_paragraph()
    p2.add_run("Ссылка: ").bold = True
    p2.add_run("https://www.inform.kz/ru/svishe-28-tisyach-zhalob-na-pereboi-s-elektrichestvom-postupilo-v-shimkente-b4549b")

    doc.add_heading("8. Telegram-бот: короткая вставка в демо", level=1)
    add_bullet(doc, "Роль в проекте: не отдельный продукт, а канал входящих обращений из чатов ЖК, дворов и инициативных групп.")
    add_bullet(doc, "Команды: `/qala_start` - подключить чат; `/qala_collect 30m` - начать сбор; `/qala_status` - проверить окно; `/qala_collect_stop` - отправить сообщения в Qala AI.")
    add_bullet(doc, "Что происходит внутри: бот сохраняет адрес/геолокацию чата, буферизует сообщения, отбрасывает шум, группирует похожие жалобы и создает заявки через внутренний API.")
    add_bullet(doc, "Как сказать за 15 секунд: 'Веб-форма - это один канал. Но жители часто пишут в Telegram-чатах ЖК. Поэтому у нас есть бот: он собирает сообщения за период, убирает шум, группирует похожие жалобы и отправляет их в Qala AI как заявки'.")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Qala AI | Smart City Hackathon | Спикерский таймлайн").font.size = Pt(8)

    doc.save(OUT)


if __name__ == "__main__":
    build()
